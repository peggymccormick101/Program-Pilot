"""Renders the structured output of ai.generate_roadmap_options() into the
executive-style Word document the workflow step promises. Claude supplies
the analysis; this module owns the formatting, so every run looks the
same regardless of what Claude returned that time.

Colors are Program Pilot's brand palette (from the marketing reference
image): dark navy for titles/table headers, royal blue for section
headings. Green/yellow stay as functional colors (recommended / caution),
independent of brand."""

import io
from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

NAVY = RGBColor(0x13, 0x2A, 0x52)
ROYAL_BLUE = RGBColor(0x1F, 0x5F, 0xCB)
GREY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

GREEN_FILL = "E2F0D9"
YELLOW_FILL = "FFF2CC"
NAVY_FILL = "132A52"
LIGHT_BLUE_FILL = "E4EDFB"


def _shade_cell(cell, hex_fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, *, bold=False, color: RGBColor = None, size=10):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text or "")
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _heading(doc, text: str, *, size: int, color: RGBColor, space_before=18, space_after=6):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return paragraph


def _callout(doc, text: str, fill_hex: str, *, bold_prefix: str = None):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _shade_cell(cell, fill_hex)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if bold_prefix:
        r = paragraph.add_run(bold_prefix + "  ")
        r.bold = True
        r.font.size = Pt(10.5)
    r2 = paragraph.add_run(text)
    r2.font.size = Pt(10.5)
    doc.add_paragraph()


def _key_value_table(doc, rows: list[dict]):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header = table.rows[0].cells
    _set_cell_text(header[0], "Planning input", bold=True, color=WHITE)
    _set_cell_text(header[1], "Treatment", bold=True, color=WHITE)
    for cell in header:
        _shade_cell(cell, NAVY_FILL)
    for row in rows:
        cells = table.add_row().cells
        _set_cell_text(cells[0], row.get("label", ""), bold=True)
        _set_cell_text(cells[1], row.get("value", ""))
    doc.add_paragraph()


def _comparison_table(doc, rows: list[dict], recommended: str):
    headers = ["Option", "Primary outcome", "Planned releases", "Relative risk", "Tradeoff"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, label, bold=True, color=WHITE)
        _shade_cell(cell, NAVY_FILL)
    for row in rows:
        cells = table.add_row().cells
        values = [
            row.get("option", ""),
            row.get("primary_outcome", ""),
            row.get("planned_releases", ""),
            row.get("relative_risk", ""),
            row.get("tradeoff", ""),
        ]
        is_recommended = row.get("option") == recommended
        for cell, value in zip(cells, values):
            _set_cell_text(cell, value, bold=is_recommended)
            if is_recommended:
                _shade_cell(cell, GREEN_FILL)
    doc.add_paragraph()


def _releases_table(doc, releases: list[dict]):
    headers = ["Release", "Features", "Frontend staff-days", "Backend staff-days", "Sequencing rationale"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, label, bold=True, color=WHITE)
        _shade_cell(cell, NAVY_FILL)
    for release in releases:
        cells = table.add_row().cells
        values = [
            release.get("release_label", ""),
            release.get("features", ""),
            release.get("frontend_days", ""),
            release.get("backend_days", ""),
            release.get("rationale", ""),
        ]
        for cell, value in zip(cells, values):
            _set_cell_text(cell, value)
    doc.add_paragraph()


def _bullets(doc, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_roadmap_docx(
    data: dict, program_name: str = "Program Pilot", release_number: str = None
) -> bytes:
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Draft Multi-Year Roadmap Options")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = NAVY

    program_label = f"{program_name} — Release {release_number}" if release_number else program_name
    subtitle = doc.add_paragraph()
    run = subtitle.add_run(f"Executive review | {program_label} | {date.today():%B %d, %Y}")
    run.font.size = Pt(12)
    run.font.color.rgb = GREY
    doc.add_paragraph()

    _callout(doc, data.get("recommendation_summary", ""), LIGHT_BLUE_FILL, bold_prefix="Recommendation")

    _heading(doc, "Planning basis", size=16, color=ROYAL_BLUE)
    _key_value_table(doc, data.get("planning_basis", []))

    _heading(doc, "Executive comparison", size=16, color=ROYAL_BLUE)
    recommended = data.get("recommended_option_name", "")
    _comparison_table(doc, data.get("executive_comparison", []), recommended)

    for index, option in enumerate(data.get("options", []), start=1):
        _heading(doc, f"Option {index} - {option.get('name', '')}", size=16, color=ROYAL_BLUE)
        doc.add_paragraph(option.get("intro", ""))
        _releases_table(doc, option.get("releases", []))

        _heading(doc, "Key rationale", size=12, color=NAVY, space_before=6)
        doc.add_paragraph(option.get("key_rationale", ""))
        _heading(doc, "Key risks", size=12, color=NAVY, space_before=6)
        doc.add_paragraph(option.get("key_risks", ""))
        _heading(doc, "Deferred features", size=12, color=NAVY, space_before=6)
        doc.add_paragraph(option.get("deferred_features", ""))

    _heading(doc, "Recommendation for PdM Review", size=16, color=ROYAL_BLUE)
    _callout(doc, f"Recommended option: {recommended}", GREEN_FILL)

    _heading(doc, "Why this option", size=12, color=NAVY, space_before=6)
    _bullets(doc, data.get("why_this_option", []))

    _heading(doc, "PdM review decisions", size=12, color=NAVY, space_before=6)
    _bullets(doc, data.get("pdm_review_decisions", []))

    doc.add_paragraph()
    _callout(doc, data.get("disclaimer", ""), YELLOW_FILL, bold_prefix="Required disclaimer")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
